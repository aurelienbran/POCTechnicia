"""
Implémentation standard du service de conversion de documents.
Prend en charge les formats de fichiers les plus courants.
"""

import logging
import asyncio
import time
import io
import os
import tempfile
from typing import List, Dict, Any, Optional, Union, BinaryIO
from pathlib import Path

from .base import DocumentConverter, ConversionResult, ConversionError

logger = logging.getLogger(__name__)

class StandardDocumentConverter(DocumentConverter):
    """
    Implémentation standard du service de conversion de documents.
    Utilise des bibliothèques Python standard pour convertir divers formats.
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None):
        """
        Initialise le convertisseur standard.
        
        Args:
            config: Configuration du convertisseur
                - temp_dir: Répertoire temporaire (défaut: tempfile.gettempdir())
                - pdf_parser: Parser PDF à utiliser ('pypdf', 'pdfminer', 'fitz') (défaut: 'fitz')
                - office_parser: Parser Office à utiliser ('python-docx', 'odf') (défaut: 'python-docx')
        """
        super().__init__(config)
        
        # Configuration
        self.temp_dir = self.config.get('temp_dir', tempfile.gettempdir())
        self.pdf_parser = self.config.get('pdf_parser', 'fitz')
        self.office_parser = self.config.get('office_parser', 'python-docx')
        
        # Dépendances
        self._dependencies = {
            'pdf': {'fitz': False, 'pypdf': False, 'pdfminer': False},
            'office': {'python-docx': False, 'odf': False},
            'image': {'pytesseract': False},
            'html': {'bs4': False}
        }
        
        # Extensions supportées par catégorie
        self._supported_extensions = {
            'pdf': ['.pdf'],
            'office': ['.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.odt', '.ods', '.odp'],
            'text': ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.htm', '.rtf'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.tif'],
        }
    
    @property
    def provider_name(self) -> str:
        """
        Nom du provider de conversion.
        
        Returns:
            "standard"
        """
        return "standard"
    
    async def initialize(self) -> bool:
        """
        Initialise le convertisseur et vérifie les dépendances.
        
        Returns:
            True si l'initialisation a réussi, False sinon
        """
        if self.initialized:
            return True
        
        try:
            # Vérifier les dépendances PDF
            if self.pdf_parser == 'fitz':
                try:
                    import fitz
                    self._dependencies['pdf']['fitz'] = True
                except ImportError:
                    logger.warning("PyMuPDF (fitz) n'est pas installé")
            
            if self.pdf_parser == 'pypdf' or not self._dependencies['pdf'][self.pdf_parser]:
                try:
                    import pypdf
                    self._dependencies['pdf']['pypdf'] = True
                except ImportError:
                    logger.warning("PyPDF2 n'est pas installé")
            
            if self.pdf_parser == 'pdfminer' or (not self._dependencies['pdf']['fitz'] and not self._dependencies['pdf']['pypdf']):
                try:
                    import pdfminer
                    self._dependencies['pdf']['pdfminer'] = True
                except ImportError:
                    logger.warning("pdfminer n'est pas installé")
            
            # Vérifier les dépendances Office
            if self.office_parser == 'python-docx':
                try:
                    import docx
                    self._dependencies['office']['python-docx'] = True
                except ImportError:
                    logger.warning("python-docx n'est pas installé")
            
            if self.office_parser == 'odf' or not self._dependencies['office'][self.office_parser]:
                try:
                    import odf
                    self._dependencies['office']['odf'] = True
                except ImportError:
                    logger.warning("odfpy n'est pas installé")
            
            # Vérifier les dépendances d'images
            try:
                import pytesseract
                self._dependencies['image']['pytesseract'] = True
            except ImportError:
                logger.warning("pytesseract n'est pas installé")
            
            # Vérifier les dépendances HTML
            try:
                import bs4
                self._dependencies['html']['bs4'] = True
            except ImportError:
                logger.warning("BeautifulSoup n'est pas installé")
            
            # Vérifier si au moins un parser est disponible pour chaque catégorie
            pdf_available = any(self._dependencies['pdf'].values())
            office_available = any(self._dependencies['office'].values())
            
            # Journaliser le statut d'initialisation
            logger.info(f"Statut d'initialisation du convertisseur standard:")
            logger.info(f"  PDF parsers: {self._dependencies['pdf']}")
            logger.info(f"  Office parsers: {self._dependencies['office']}")
            logger.info(f"  Image parsers: {self._dependencies['image']}")
            logger.info(f"  HTML parsers: {self._dependencies['html']}")
            
            # Considérer comme réussi si au moins les parsers PDF et texte sont disponibles
            self.initialized = pdf_available
            return self.initialized
            
        except Exception as e:
            logger.error(f"Erreur lors de l'initialisation du convertisseur standard: {str(e)}")
            return False
    
    async def convert_file(self, 
                    file_path: Union[str, Path],
                    output_format: str = "text",
                    **kwargs) -> ConversionResult:
        """
        Convertit un fichier en texte.
        
        Args:
            file_path: Chemin vers le fichier à convertir
            output_format: Format de sortie (défaut: "text")
            **kwargs: Options spécifiques au convertisseur
                - force_ocr: Forcer l'OCR pour les PDFs (défaut: False)
                - extract_metadata: Extraire les métadonnées (défaut: True)
                - dpi: Résolution pour l'OCR (défaut: 300)
                
        Returns:
            Résultat de la conversion
        """
        start_time = time.time()
        
        # S'assurer que le convertisseur est initialisé
        if not self.initialized and not await self.initialize():
            return ConversionResult(
                success=False,
                error_message="Le convertisseur n'a pas pu être initialisé"
            )
        
        # Convertir en Path
        file_path = Path(file_path)
        
        # Vérifier l'existence du fichier
        if not file_path.exists():
            return ConversionResult(
                success=False,
                error_message=f"Le fichier {file_path} n'existe pas"
            )
        
        # Détecter le type de fichier
        mime_type = await self.detect_file_type(file_path)
        extension = file_path.suffix.lower()
        
        # Extraire les métadonnées si demandé
        metadata = {}
        if kwargs.get('extract_metadata', True):
            try:
                metadata = await self.extract_metadata(file_path)
            except Exception as e:
                logger.warning(f"Erreur lors de l'extraction des métadonnées: {str(e)}")
        
        try:
            # Convertir selon le type de fichier
            if extension in self._supported_extensions['pdf']:
                result = await self._convert_pdf(file_path, **kwargs)
            elif extension in self._supported_extensions['office']:
                result = await self._convert_office(file_path, **kwargs)
            elif extension in self._supported_extensions['text']:
                result = await self._convert_text(file_path, **kwargs)
            elif extension in self._supported_extensions['image']:
                result = await self._convert_image(file_path, **kwargs)
            else:
                return ConversionResult(
                    success=False,
                    error_message=f"Type de fichier non supporté: {extension} ({mime_type})"
                )
            
            # Ajouter les métadonnées
            if result.success and metadata:
                if result.metadata is None:
                    result.metadata = {}
                result.metadata.update(metadata)
            
            # Ajouter le temps de traitement
            result.processing_time = time.time() - start_time
            
            return result
            
        except Exception as e:
            logger.error(f"Erreur lors de la conversion du fichier {file_path}: {str(e)}")
            return ConversionResult(
                success=False,
                error_message=str(e),
                processing_time=time.time() - start_time
            )

    async def convert_bytes(self,
                     content: Union[bytes, BinaryIO],
                     file_type: str,
                     output_format: str = "text",
                     **kwargs) -> ConversionResult:
        """
        Convertit un contenu binaire en texte.
        
        Args:
            content: Contenu binaire à convertir
            file_type: Type du fichier (extension ou MIME type)
            output_format: Format de sortie (défaut: "text")
            **kwargs: Options spécifiques au convertisseur
            
        Returns:
            Résultat de la conversion
        """
        start_time = time.time()
        
        # S'assurer que le convertisseur est initialisé
        if not self.initialized and not await self.initialize():
            return ConversionResult(
                success=False,
                error_message="Le convertisseur n'a pas pu être initialisé"
            )
        
        # Créer un fichier temporaire
        temp_dir = Path(self.temp_dir)
        temp_dir.mkdir(parents=True, exist_ok=True)
        
        # Déterminer l'extension
        if file_type.startswith('.'):
            extension = file_type
        elif '/' in file_type:  # MIME type
            extension = next((ext for ext, mime in self._get_mime_types().items() if mime == file_type), '.bin')
        else:
            extension = f".{file_type}"
        
        try:
            # Écrire le contenu dans un fichier temporaire
            with tempfile.NamedTemporaryFile(suffix=extension, delete=False, dir=temp_dir) as temp_file:
                temp_path = temp_file.name
                
                if isinstance(content, bytes):
                    temp_file.write(content)
                else:
                    temp_file.write(content.read())
            
            # Convertir le fichier temporaire
            result = await self.convert_file(temp_path, output_format, **kwargs)
            
            return result
            
        except Exception as e:
            logger.error(f"Erreur lors de la conversion du contenu binaire: {str(e)}")
            return ConversionResult(
                success=False,
                error_message=str(e),
                processing_time=time.time() - start_time
            )
        finally:
            # Supprimer le fichier temporaire
            if 'temp_path' in locals() and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except Exception as e:
                    logger.warning(f"Erreur lors de la suppression du fichier temporaire: {str(e)}")
    
    async def extract_metadata(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Extrait les métadonnées d'un fichier.
        
        Args:
            file_path: Chemin vers le fichier
            
        Returns:
            Dictionnaire des métadonnées
        """
        # Convertir en Path
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # Métadonnées de base (taille, date, etc.)
        metadata = {
            "filename": file_path.name,
            "extension": extension,
            "size_bytes": file_path.stat().st_size,
            "modified_time": file_path.stat().st_mtime,
        }
        
        # Extraire des métadonnées spécifiques selon le type de fichier
        try:
            if extension in self._supported_extensions['pdf']:
                pdf_metadata = await self._extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            elif extension in self._supported_extensions['office']:
                office_metadata = await self._extract_office_metadata(file_path)
                metadata.update(office_metadata)
            # Ajouter d'autres types de fichiers au besoin
        except Exception as e:
            logger.warning(f"Erreur lors de l'extraction des métadonnées spécifiques: {str(e)}")
        
        return metadata
    
    async def supported_file_types(self) -> List[str]:
        """
        Retourne la liste des types de fichiers supportés.
        
        Returns:
            Liste des extensions de fichiers supportées
        """
        supported = []
        
        # PDF (si disponible)
        if any(self._dependencies['pdf'].values()):
            supported.extend(self._supported_extensions['pdf'])
        
        # Office (si disponible)
        if any(self._dependencies['office'].values()):
            supported.extend(self._supported_extensions['office'])
        
        # Texte (toujours supporté)
        supported.extend(self._supported_extensions['text'])
        
        # Images (si pytesseract est disponible)
        if self._dependencies['image']['pytesseract']:
            supported.extend(self._supported_extensions['image'])
        
        return supported
    
    def _get_mime_types(self) -> Dict[str, str]:
        """
        Retourne un dictionnaire des types MIME associés aux extensions.
        
        Returns:
            Dictionnaire {extension: mime_type}
        """
        import mimetypes
        
        # S'assurer que les types MIME sont initialisés
        if not mimetypes.inited:
            mimetypes.init()
            
        # Créer un dictionnaire extension -> type MIME
        mime_types = {}
        for ext in [ext for exts in self._supported_extensions.values() for ext in exts]:
            mime_type, _ = mimetypes.guess_type(f"file{ext}")
            if mime_type:
                mime_types[ext] = mime_type
        
        return mime_types
    
    # Méthodes d'extraction de métadonnées spécifiques aux formats
    
    async def _extract_pdf_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extrait les métadonnées d'un fichier PDF.
        
        Args:
            file_path: Chemin vers le fichier PDF
            
        Returns:
            Dictionnaire des métadonnées
        """
        metadata = {}
        
        # Utiliser PyMuPDF si disponible
        if self._dependencies['pdf']['fitz']:
            try:
                import fitz
                doc = fitz.open(str(file_path))
                
                metadata.update({
                    "page_count": doc.page_count,
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "subject": doc.metadata.get("subject", ""),
                    "keywords": doc.metadata.get("keywords", ""),
                    "creator": doc.metadata.get("creator", ""),
                    "producer": doc.metadata.get("producer", ""),
                })
                
                # Taille du document
                if doc.page_count > 0:
                    first_page = doc.load_page(0)
                    metadata["page_size"] = {"width": first_page.rect.width, "height": first_page.rect.height}
                
                # Autres métadonnées intéressantes
                metadata["has_toc"] = doc.get_toc().count > 0
                
                doc.close()
                return metadata
                
            except Exception as e:
                logger.warning(f"Erreur lors de l'extraction des métadonnées PDF avec PyMuPDF: {str(e)}")
        
        # Fallback: utiliser PyPDF2
        if self._dependencies['pdf']['pypdf']:
            try:
                import pypdf
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    
                    document_info = reader.metadata or {}
                    
                    metadata.update({
                        "page_count": len(reader.pages),
                        "title": document_info.get("/Title", ""),
                        "author": document_info.get("/Author", ""),
                        "subject": document_info.get("/Subject", ""),
                        "keywords": document_info.get("/Keywords", ""),
                        "creator": document_info.get("/Creator", ""),
                        "producer": document_info.get("/Producer", ""),
                    })
                    
                    return metadata
                    
            except Exception as e:
                logger.warning(f"Erreur lors de l'extraction des métadonnées PDF avec PyPDF2: {str(e)}")
        
        return metadata
    
    async def _extract_office_metadata(self, file_path: Path) -> Dict[str, Any]:
        """
        Extrait les métadonnées d'un fichier Office.
        
        Args:
            file_path: Chemin vers le fichier Office
            
        Returns:
            Dictionnaire des métadonnées
        """
        metadata = {}
        extension = file_path.suffix.lower()
        
        # Word (DOCX)
        if extension == '.docx' and self._dependencies['office']['python-docx']:
            try:
                import docx
                doc = docx.Document(file_path)
                
                # Propriétés du document
                core_properties = doc.core_properties
                
                metadata.update({
                    "title": core_properties.title or "",
                    "author": core_properties.author or "",
                    "subject": core_properties.subject or "",
                    "keywords": core_properties.keywords or "",
                    "created": core_properties.created,
                    "modified": core_properties.modified,
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                })
                
                return metadata
                
            except Exception as e:
                logger.warning(f"Erreur lors de l'extraction des métadonnées DOCX: {str(e)}")
        
        # TODO: Ajouter la prise en charge d'autres formats Office
        
        return metadata
        
    # Méthodes de conversion spécifiques aux formats
    
    async def _convert_pdf(self, file_path: Path, **kwargs) -> ConversionResult:
        """
        Convertit un fichier PDF en texte.
        
        Args:
            file_path: Chemin vers le fichier PDF
            **kwargs: Options spécifiques
                - force_ocr: Forcer l'OCR (défaut: False)
                - dpi: Résolution pour l'OCR (défaut: 300)
                
        Returns:
            Résultat de la conversion
        """
        # Une implémentation simplifiée qui utilise PyMuPDF (fitz)
        if self._dependencies['pdf']['fitz']:
            try:
                import fitz
                doc = fitz.open(str(file_path))
                
                text_content = ""
                for page_num in range(doc.page_count):
                    page = doc.load_page(page_num)
                    text_content += page.get_text()
                
                doc.close()
                
                return ConversionResult(
                    success=True,
                    text_content=text_content,
                    pages_processed=doc.page_count,
                    total_pages=doc.page_count,
                    metadata={"parser": "fitz"}
                )
                
            except Exception as e:
                logger.error(f"Erreur lors de la conversion PDF avec PyMuPDF: {str(e)}")
        
        # Fallback: utiliser PyPDF2
        if self._dependencies['pdf']['pypdf']:
            try:
                import pypdf
                with open(file_path, 'rb') as f:
                    reader = pypdf.PdfReader(f)
                    
                    text_content = ""
                    for page in reader.pages:
                        text_content += page.extract_text() or ""
                    
                    return ConversionResult(
                        success=True,
                        text_content=text_content,
                        pages_processed=len(reader.pages),
                        total_pages=len(reader.pages),
                        metadata={"parser": "pypdf"}
                    )
                    
            except Exception as e:
                logger.error(f"Erreur lors de la conversion PDF avec PyPDF2: {str(e)}")
        
        # Si aucun parser n'a fonctionné
        raise ConversionError("Aucun parser PDF disponible n'a pu convertir le fichier")
    
    async def _convert_text(self, file_path: Path, **kwargs) -> ConversionResult:
        """
        Convertit un fichier texte.
        
        Args:
            file_path: Chemin vers le fichier texte
            **kwargs: Options spécifiques
                - encoding: Encodage du fichier (défaut: 'utf-8')
                
        Returns:
            Résultat de la conversion
        """
        encoding = kwargs.get('encoding', 'utf-8')
        extension = file_path.suffix.lower()
        
        # Fichiers texte simples
        if extension in ['.txt', '.md', '.csv']:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                
                return ConversionResult(
                    success=True,
                    text_content=text_content,
                    pages_processed=1,
                    total_pages=1,
                    metadata={"format": extension[1:]}
                )
                
            except UnicodeDecodeError:
                # Essayer avec un autre encodage
                for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        with open(file_path, 'r', encoding=enc) as f:
                            text_content = f.read()
                        
                        return ConversionResult(
                            success=True,
                            text_content=text_content,
                            pages_processed=1,
                            total_pages=1,
                            metadata={"format": extension[1:], "encoding": enc}
                        )
                    except UnicodeDecodeError:
                        continue
                
                # Si aucun encodage n'a fonctionné
                raise ConversionError(f"Impossible de décoder le fichier avec les encodages essayés")
        
        # Fichiers HTML
        elif extension in ['.html', '.htm']:
            if self._dependencies['html']['bs4']:
                try:
                    from bs4 import BeautifulSoup
                    
                    with open(file_path, 'r', encoding=encoding) as f:
                        soup = BeautifulSoup(f.read(), 'html.parser')
                        
                        # Supprimer les scripts et styles
                        for script in soup(["script", "style"]):
                            script.extract()
                        
                        text_content = soup.get_text(separator="\n")
                        
                        # Nettoyer le texte (espaces multiples, lignes vides)
                        text_content = '\n'.join([line.strip() for line in text_content.split('\n') if line.strip()])
                        
                        return ConversionResult(
                            success=True,
                            text_content=text_content,
                            pages_processed=1,
                            total_pages=1,
                            metadata={"format": "html"}
                        )
                except Exception as e:
                    logger.error(f"Erreur lors de la conversion HTML: {str(e)}")
            
            # Fallback: lire comme texte simple
            return await self._convert_text(file_path, encoding=encoding)
        
        # Autres formats texte
        else:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text_content = f.read()
                
                return ConversionResult(
                    success=True,
                    text_content=text_content,
                    pages_processed=1,
                    total_pages=1,
                    metadata={"format": extension[1:]}
                )
            except Exception as e:
                logger.error(f"Erreur lors de la conversion du fichier texte {file_path}: {str(e)}")
                raise ConversionError(f"Impossible de convertir le fichier texte: {str(e)}")
    
    async def _convert_office(self, file_path: Path, **kwargs) -> ConversionResult:
        """
        Convertit un fichier Office en texte.
        
        Args:
            file_path: Chemin vers le fichier Office
            **kwargs: Options spécifiques
                
        Returns:
            Résultat de la conversion
        """
        extension = file_path.suffix.lower()
        
        # DOCX
        if extension == '.docx' and self._dependencies['office']['python-docx']:
            try:
                import docx
                doc = docx.Document(file_path)
                
                # Extraire le texte des paragraphes
                paragraphs = [p.text for p in doc.paragraphs]
                
                # Extraire le texte des tableaux
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text for cell in row.cells])
                        if row_text.strip():
                            paragraphs.append(row_text)
                
                text_content = '\n\n'.join(paragraphs)
                
                return ConversionResult(
                    success=True,
                    text_content=text_content,
                    pages_processed=1,  # On ne peut pas facilement compter les pages dans un DOCX
                    total_pages=1,
                    metadata={
                        "format": "docx",
                        "paragraph_count": len(doc.paragraphs),
                        "table_count": len(doc.tables)
                    }
                )
                
            except Exception as e:
                logger.error(f"Erreur lors de la conversion DOCX: {str(e)}")
        
        # TODO: Ajouter la prise en charge d'autres formats Office
        
        # Si aucun parser n'a fonctionné
        raise ConversionError(f"Conversion du format {extension} non prise en charge")
    
    async def _convert_image(self, file_path: Path, **kwargs) -> ConversionResult:
        """
        Convertit une image en texte via OCR.
        
        Args:
            file_path: Chemin vers l'image
            **kwargs: Options spécifiques
                - lang: Langue pour l'OCR (défaut: 'fra')
                - dpi: Résolution pour l'OCR (défaut: 300)
                
        Returns:
            Résultat de la conversion
        """
        if not self._dependencies['image']['pytesseract']:
            raise ConversionError("pytesseract n'est pas disponible pour l'OCR d'images")
        
        try:
            import pytesseract
            from PIL import Image
            
            # Options OCR
            lang = kwargs.get('lang', 'fra')  # fra pour français
            dpi = kwargs.get('dpi', 300)
            
            # Ouvrir l'image
            img = Image.open(file_path)
            
            # Effectuer l'OCR
            text_content = pytesseract.image_to_string(img, lang=lang)
            
            return ConversionResult(
                success=True,
                text_content=text_content,
                pages_processed=1,
                total_pages=1,
                metadata={
                    "format": file_path.suffix[1:],
                    "ocr_lang": lang,
                    "image_size": {"width": img.width, "height": img.height}
                }
            )
            
        except Exception as e:
            logger.error(f"Erreur lors de l'OCR de l'image: {str(e)}")
            raise ConversionError(f"Impossible d'extraire le texte de l'image: {str(e)}")
